"""Ch14-05: Export to DOCX - a formal Word report with python-docx.

Bilingual scene (zh/en) driven by ``EpiBaseScene.t()``. All on-screen prose is
read from ``TEXT`` via ``self.t(key)``; code strings stay identical (ASCII) across
languages. The generated document is drawn as a mockup from ``RoundedRectangle`` /
``Rectangle`` / ``Line`` / ``Text`` (no new mobjects).
"""

from __future__ import annotations

from manim import (
    DOWN,
    LEFT,
    ORIGIN,
    RIGHT,
    UP,
    FadeIn,
    FadeOut,
    Line,
    ManimColor,
    Rectangle,
    RoundedRectangle,
    Text,
    VGroup,
)

from videos.src.base_scene import EpiBaseScene
from videos.src.code_mobjects import (
    ACCENT_ORANGE,
    BG_CARD,
    BG_CARD_ALT,
    BORDER_LIGHT,
    FONT_CJK,
    TEXT_PRIMARY,
    TEXT_SECONDARY,
    BlindSpotBanner,
    ExtraExampleBanner,
)


class Ch14ExportDocxScene(EpiBaseScene):
    """Tutorial video scene: exporting a formal DOCX report with python-docx."""

    total_steps: int = 9

    TEXT: dict[str, dict[str, str]] = {
        "zh": {
            "title_main": "輸出正式 DOCX 報告",
            "title_sub": "python-docx：標題、段落、內嵌圖",
            "why_heading": "為什麼還要 DOCX？",
            "why_lines": [
                "簡報給眼睛掃，報告給人逐字讀",
                "正式標題階層、段落、表格樣式",
                "適合當結案文件歸檔、簽核",
                "→ 跟 PPTX 一樣可重現",
            ],
            "docx_code_heading": "python-docx：標題 + 段落 + 內嵌圖",
            "docx_code_title": "export_docx.py",
            "doc_mockup_heading": "產出的報告長這樣",
            "doc_title": "退伍軍人症群聚調查報告",
            "doc_h1_summary": "摘要",
            "doc_h1_figure": "流行曲線",
            "doc_fig_caption": "流行曲線（內嵌圖）",
            "reproducible_heading": "呼應 Ch13：可重現",
            "reproducible_lines": [
                "報告不是手打的，是程式長出來的",
                "資料更新 → 重跑一次，報告重生",
                "永遠沒有「這用哪一版的圖」的疑問",
                "→ 結果永遠跟分析同步",
            ],
            "summary_heading": "DOCX 輸出三重點",
            "summary_lines": [
                "① add_heading 的 level 控制標題階層",
                "② add_paragraph 內文、add_picture 內嵌圖",
                "③ 與 PPTX 共用素材，換工具不換資料",
                "→ 一鍵產出正式結案報告",
            ],
            "extra_banner_title": "額外範例：諾羅病毒院內群聚正式報告",
            "extra_formal_heading": "迴圈跑八段落，自動生報告",
            "extra_formal_title": "norovirus_report.py",
            "blindspot_banner_title": "DOCX 輸出三個新手地雷",
            "outro_heading": "全系列完結：你已能獨當一面",
            "outro_sub": "從接獲通報到一鍵輸出報告，全套走完",
        },
        "en": {
            "title_main": "Export a Formal DOCX Report",
            "title_sub": "python-docx: headings, paragraphs, embedded figures",
            "why_heading": "Why Also a DOCX?",
            "why_lines": [
                "A deck is scanned; a report is read word by word",
                "Formal heading levels, paragraphs, table styles",
                "Fit to archive and sign off as the closing document",
                "-> reproducible, just like the PPTX",
            ],
            "docx_code_heading": "python-docx: heading + paragraph + figure",
            "docx_code_title": "export_docx.py",
            "doc_mockup_heading": "What the Report Looks Like",
            "doc_title": "Legionnaires' cluster investigation report",
            "doc_h1_summary": "Summary",
            "doc_h1_figure": "Epidemic curve",
            "doc_fig_caption": "Epi curve (embedded figure)",
            "reproducible_heading": "Ch13 echoed: reproducible",
            "reproducible_lines": [
                "The report isn't typed by hand - code grows it",
                "Data updates -> rerun once, the report rebuilds",
                'Never "which version of the chart was this?"',
                "-> results always stay in sync with the analysis",
            ],
            "summary_heading": "Three Takeaways on DOCX Export",
            "summary_lines": [
                "1. add_heading level controls the heading hierarchy",
                "2. add_paragraph for body, add_picture to embed a figure",
                "3. Shares assets with PPTX - swap tool, keep the data",
                "-> a formal closing report in one click",
            ],
            "extra_banner_title": "Extra example: a formal norovirus outbreak report",
            "extra_formal_heading": "Loop the eight sections, auto-build the report",
            "extra_formal_title": "norovirus_report.py",
            "blindspot_banner_title": "Three Beginner Blind Spots",
            "outro_heading": "Series finale: you can stand on your own",
            "outro_sub": "From notification to one-click report - the full workflow",
        },
    }

    def construct(self) -> None:
        self.construct_from_segments()

    # ------------------------------------------------------------------
    # Shared helpers
    # ------------------------------------------------------------------

    def _bullets(self, heading_key: str, lines_key: str, duration: float) -> None:
        heading = self.t(heading_key)
        lines = self.t(lines_key)
        h = Text(heading, font=FONT_CJK, font_size=30, color=ACCENT_ORANGE).to_edge(UP, buff=0.8)
        bl = (
            VGroup(*[Text(x, font=FONT_CJK, font_size=22, color=TEXT_PRIMARY) for x in lines])
            .arrange(DOWN, aligned_edge=LEFT, buff=0.4)
            .next_to(h, DOWN, buff=0.55)
        )
        if bl.width > 12.5:
            bl.scale_to_fit_width(12.5)
        self.play(FadeIn(h), run_time=0.5)
        self.play(FadeIn(bl, lag_ratio=0.2), run_time=1.2)
        self.wait(max(0.1, duration - 1.7))
        self.play(FadeOut(VGroup(h, bl)), run_time=0.5)

    def _code_block(self, heading_key: str, title_key: str, code: str, duration: float) -> None:
        h = Text(self.t(heading_key), font=FONT_CJK, font_size=26, color=ACCENT_ORANGE).to_edge(
            UP, buff=0.5
        )
        self.play(FadeIn(h), run_time=0.4)
        panel = self.show_code(code, title=self.t(title_key), position=ORIGIN + DOWN * 0.3)
        self.wait(max(0.1, duration - 1.4))
        self.play(FadeOut(VGroup(h, panel)), run_time=0.5)

    # ------------------------------------------------------------------
    # Main lesson
    # ------------------------------------------------------------------

    def show_title(self, duration: float = 3.0, **kwargs) -> None:
        self.show_title_card(self.t("title_main"), self.t("title_sub"), duration=duration)

    def show_why_docx(self, duration: float = 7.0, **kwargs) -> None:
        self.show_step_indicator(1, self.total_steps)
        self._bullets("why_heading", "why_lines", duration)

    def show_docx_code(self, duration: float = 9.0, **kwargs) -> None:
        self.show_step_indicator(2, self.total_steps)
        code = kwargs.get(
            "code",
            (
                "from docx import Document\n"
                "from docx.shared import Inches as DInches\n"
                "doc = Document()\n"
                'doc.add_heading("Legionella cluster report", level=0)\n'
                'doc.add_heading("Summary", level=1)\n'
                'doc.add_paragraph(f"AR {ar:.1%}, CFR {cfr:.1%}")\n'
                "doc.add_picture(buf, width=DInches(5.5))\n"
                'doc.save("legionella_report.docx")'
            ),
        )
        self._code_block("docx_code_heading", "docx_code_title", code, duration)

    def show_doc_mockup(self, duration: float = 8.0, **kwargs) -> None:
        self.show_step_indicator(3, self.total_steps)

        heading = Text(
            self.t("doc_mockup_heading"), font=FONT_CJK, font_size=28, color=ACCENT_ORANGE
        ).to_edge(UP, buff=0.5)

        page = RoundedRectangle(
            corner_radius=0.12,
            width=6.8,
            height=6.2,
            fill_color=ManimColor(BG_CARD),
            fill_opacity=1,
            stroke_color=ManimColor(BORDER_LIGHT),
            stroke_width=2,
        )
        title = Text(
            self.t("doc_title"), font=FONT_CJK, font_size=20, color=TEXT_PRIMARY, weight="BOLD"
        )
        if title.width > 5.6:
            title.scale_to_fit_width(5.6)
        rule = Line(ORIGIN, RIGHT * 5.6, color=ManimColor(BORDER_LIGHT), stroke_width=2)
        h1a = Text(self.t("doc_h1_summary"), font=FONT_CJK, font_size=17, color=ACCENT_ORANGE)
        para = VGroup(
            Line(ORIGIN, RIGHT * 5.4, color=ManimColor(BORDER_LIGHT), stroke_width=4),
            Line(ORIGIN, RIGHT * 5.4, color=ManimColor(BORDER_LIGHT), stroke_width=4),
            Line(ORIGIN, RIGHT * 3.6, color=ManimColor(BORDER_LIGHT), stroke_width=4),
        ).arrange(DOWN, aligned_edge=LEFT, buff=0.22)
        h1b = Text(self.t("doc_h1_figure"), font=FONT_CJK, font_size=17, color=ACCENT_ORANGE)
        figbox = Rectangle(
            width=4.8,
            height=1.6,
            fill_color=ManimColor(BG_CARD_ALT),
            fill_opacity=1,
            stroke_color=ManimColor(BORDER_LIGHT),
            stroke_width=1.5,
        )
        figcap = Text(
            self.t("doc_fig_caption"), font=FONT_CJK, font_size=14, color=TEXT_SECONDARY
        ).move_to(figbox.get_center())
        figure = VGroup(figbox, figcap)

        content = VGroup(title, rule, h1a, para, h1b, figure).arrange(
            DOWN, aligned_edge=LEFT, buff=0.3
        )
        content.move_to(page.get_center())
        doc = VGroup(page, content).next_to(heading, DOWN, buff=0.3)

        self.play(FadeIn(heading), run_time=0.4)
        self.play(FadeIn(page), run_time=0.5)
        self.play(FadeIn(content, lag_ratio=0.15), run_time=1.6)
        self.wait(max(0.1, duration - 2.9))
        self.play(FadeOut(VGroup(heading, doc)), run_time=0.5)

    def show_reproducible(self, duration: float = 7.0, **kwargs) -> None:
        self.show_step_indicator(4, self.total_steps)
        self._bullets("reproducible_heading", "reproducible_lines", duration)

    def show_main_summary(self, duration: float = 6.0, **kwargs) -> None:
        self.show_step_indicator(5, self.total_steps)
        self._bullets("summary_heading", "summary_lines", duration)

    # ------------------------------------------------------------------
    # Extra example
    # ------------------------------------------------------------------

    def show_extra_banner(self, duration: float = 2.0, **kwargs) -> None:
        self.show_section_banner(
            ExtraExampleBanner(self.t("extra_banner_title")), duration=duration
        )

    def show_extra_formal_report(self, duration: float = 8.0, **kwargs) -> None:
        self.show_step_indicator(6, self.total_steps)
        code = kwargs.get(
            "code",
            (
                "doc = Document()\n"
                'doc.add_heading("Norovirus outbreak report", level=0)\n'
                "for sec in REPORT_8_SECTIONS:\n"
                "    doc.add_heading(sec.title, level=1)\n"
                "    doc.add_paragraph(sec.body)\n"
                "doc.add_picture(epi_curve, width=DInches(6))\n"
                'doc.save("norovirus_report.docx")'
            ),
        )
        self._code_block("extra_formal_heading", "extra_formal_title", code, duration)

    # ------------------------------------------------------------------
    # Blind spots
    # ------------------------------------------------------------------

    def show_blindspot_banner(self, duration: float = 2.0, **kwargs) -> None:
        self.show_section_banner(
            BlindSpotBanner(self.t("blindspot_banner_title")), duration=duration
        )

    def show_blindspot_forgot_seek(self, duration: float = 6.0, **kwargs) -> None:
        self.show_step_indicator(7, self.total_steps)
        panel = self.show_error_vs_correct(
            kwargs.get("error_code", "doc.add_picture(buf)"),
            kwargs.get("correct_code", "buf.seek(0); doc.add_picture(buf)"),
            duration=duration,
        )
        self.play(FadeOut(panel), run_time=0.5)

    def show_blindspot_inches_clash(self, duration: float = 6.0, **kwargs) -> None:
        self.show_step_indicator(8, self.total_steps)
        panel = self.show_error_vs_correct(
            kwargs.get("error_code", "from pptx.util import Inches"),
            kwargs.get("correct_code", "from docx.shared import Inches as D"),
            duration=duration,
        )
        self.play(FadeOut(panel), run_time=0.5)

    def show_blindspot_heading_as_paragraph(self, duration: float = 6.0, **kwargs) -> None:
        self.show_step_indicator(9, self.total_steps)
        panel = self.show_error_vs_correct(
            kwargs.get("error_code", 'doc.add_paragraph("Summary")  # as title'),
            kwargs.get("correct_code", 'doc.add_heading("Summary", level=1)'),
            duration=duration,
        )
        self.play(FadeOut(panel), run_time=0.5)

    # ------------------------------------------------------------------
    # Outro
    # ------------------------------------------------------------------

    def show_outro(self, duration: float = 3.0, **kwargs) -> None:
        self.show_step_indicator(self.total_steps, self.total_steps)
        h = Text(self.t("outro_heading"), font=FONT_CJK, font_size=26, color=ACCENT_ORANGE).move_to(
            ORIGIN + UP * 0.5
        )
        s = Text(self.t("outro_sub"), font=FONT_CJK, font_size=20, color=TEXT_SECONDARY).next_to(
            h, DOWN, buff=0.4
        )
        self.play(FadeIn(h), run_time=0.6)
        self.play(FadeIn(s), run_time=0.5)
        self.wait(max(0.1, duration - 1.1))
        self.play(FadeOut(VGroup(h, s)), run_time=0.5)
